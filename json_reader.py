import json
import os
from docx import Document
from docx.shared import Pt
import re
import sys
from colorama import init
from pyfiglet import figlet_format
from termcolor import cprint
from terminaltables import AsciiTable

init(strip=not sys.stdout.isatty())  # strip colors if stdout is redirected

def main():
    cprint(figlet_format("John's Basic Reading"), 'red', attrs=['bold'])
    cprint("Hello there!\n\t The purpose of this program is to take the John's Basic Reading pdf and parse the data within it.\n"
           "as an output I will provide a set of files containing the site words, charted words, and passages all in docx files.\n"
           "In the event you are unfamiliar with this file type, this means you will be able to open it in MS word.\n"
           "All I need from you is a font size. On the following line, please type a font size, and immediately hit return.\n"
           "Alternatively you can give me a list of numbers separated by commas. \n\n"
           "For example, a single font size would look like this\n\t>> 24\n"
           "For example, multiple font sizes would look like this\n\t>> 12,14,18",
        'red')

    size_string = input("\n\t>>")
    cprint("\nPerfect! I will generate your files using size " + size_string + " font.", 'red')

    size_array = []
    dirty = size_string.split(",")
    for dirt in dirty:
        clean = dirt.strip()
        size_array.append(int(clean))


    print("--Getting started--")

    print("--Making some pretty folders--")
    os.makedirs("output")
    os.makedirs("output/form_a")
    os.makedirs("output/form_a/sight_words")
    os.makedirs("output/form_a/reading_charts")
    os.makedirs("output/form_b")
    os.makedirs("output/form_b/sight_words")
    os.makedirs("output/form_b/reading_charts")
    os.makedirs("output/form_c")
    os.makedirs("output/form_c/sight_words")
    os.makedirs("output/form_c/reading_charts")
    os.makedirs("output/form_d")
    os.makedirs("output/form_d/reading_charts")
    os.makedirs("output/form_e")
    os.makedirs("output/form_e/reading_charts")
    os.makedirs("output/ln_passages")
    os.makedirs("output/le_passages")

    print("--Chewing threw the pdf file--")
    json_data = {}
    with open('readingjson.json', encoding='utf-8') as json_data:
        json_data = json.load(json_data)

    # pull all the top level form data
    form_a = json_data.get("form_a")
    form_b = json_data.get("form_b")
    form_c = json_data.get("form_c")
    form_d = json_data.get("form_d")
    form_e = json_data.get("form_e")

    # pull all passage data
    ln_passages = json_data.get("ln_passages")
    le_passages = json_data.get("le_passages")

    print("--Working on the site words--")
    write_sight_words(form_a.get("sight_words"), "output/form_a/sight_words/", size_array)
    write_sight_words(form_b.get("sight_words"), "output/form_b/sight_words/", size_array)
    write_sight_words(form_c.get("sight_words"), "output/form_c/sight_words/", size_array)
    print("--Done with the site words--")

    print("--Working on the reading charts--")
    write_reading_chart(form_a.get("reading_chart"), "output/form_a/reading_charts/", size_array)
    write_reading_chart(form_b.get("reading_chart"), "output/form_b/reading_charts/", size_array)
    write_reading_chart(form_c.get("reading_chart"), "output/form_c/reading_charts/", size_array)
    write_reading_chart(form_d.get("reading_chart"), "output/form_d/reading_charts/", size_array)
    write_reading_chart(form_e.get("reading_chart"), "output/form_e/reading_charts/", size_array)
    print("--Done with the reading charts--")

    print("--Working on the passages--")
    write_passages(ln_passages, "output/ln_passages/", size_array)
    write_passages(le_passages, "output/le_passages/", size_array)
    print("--Done with the reading charts--")

    print("\n\n--All done! enjoy your files!--")

def write_reading_chart(reading_chart, dir_path, sizes):
    for size in sizes:
        pre_primer_1 = reading_chart.get("pre_primer_1")
        pre_primer_1_title = reading_chart.get("pre_primer_1_title")
        write_chart(pre_primer_1, pre_primer_1_title, dir_path, size, "pre_primer_1")

        pre_primer_2 = reading_chart.get("pre_primer_2")
        pre_primer_2_title = reading_chart.get("pre_primer_2_title")
        write_chart(pre_primer_2, pre_primer_2_title, dir_path, size, "pre_primer_2")

        primer = reading_chart.get("primer")
        primer_title = reading_chart.get("primer_title")
        write_chart(primer, primer_title, dir_path, size, "primer")

        grade_1 = reading_chart.get("grade_1")
        grade_1_title = reading_chart.get("grade_1_title")
        write_chart(grade_1, grade_1_title, dir_path, size, "grade_1")

        grade_2 = reading_chart.get("grade_2")
        grade_2_title = reading_chart.get("grade_2_title")
        write_chart(grade_2, grade_2_title, dir_path, size, "grade_2")

        grade_3 = reading_chart.get("grade_3")
        grade_3_title = reading_chart.get("grade_3_title")
        write_chart(grade_3, grade_3_title, dir_path, size, "grade_3")

        grade_4 = reading_chart.get("grade_4")
        grade_4_title = reading_chart.get("grade_4_title")
        write_chart(grade_4, grade_4_title, dir_path, size, "grade_4")

        grade_5 = reading_chart.get("grade_5")
        grade_5_title = reading_chart.get("grade_5_title")
        write_chart(grade_5, grade_5_title, dir_path, size, "grade_5")

        grade_6 = reading_chart.get("grade_6")
        grade_6_title = reading_chart.get("grade_6_title")
        write_chart(grade_6, grade_6_title, dir_path, size, "grade_6")

        grade_7 = reading_chart.get("grade_7")
        grade_7_title = reading_chart.get("grade_7_title")
        write_chart(grade_7, grade_7_title, dir_path, size, "grade_7")

        grade_8 = reading_chart.get("grade_8")
        grade_8_title = reading_chart.get("grade_8_title")
        write_chart(grade_8, grade_8_title, dir_path, size, "grade_8")

def write_chart(text_strings, title, dir_path, size, grade):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(size)


    p = document.add_paragraph(style='Normal')
    runner = p.add_run(title + "\n\n")

    for ts in text_strings:
        if "<b>" in ts or "<!b>" in ts:

            nb_start = re.search('(.*)<b>', ts)
            b = re.search('<b>(.*)<!b>', ts)
            nb_end = re.search('<!b>(.*)', ts)

            if nb_start:
                runner_nb_1 = p.add_run(nb_start.group(1))
            if b:
                runner_b = p.add_run(b.group(1))
                runner_b.bold = True
            if nb_end:
                runner_nb_2 = p.add_run(nb_end.group(1))

            runner_nl = p.add_run("\n")

        else:
            runner_loop = p.add_run(ts + "\n")

    sub_dir = dir_path + grade
    if not os.path.exists(sub_dir):
        os.makedirs(sub_dir)
    document.save((sub_dir + "/" + grade + "_size" + str(size) + ".docx"))

def write_passages(passages, dir_path, sizes):
    for passage in passages:
        for size in sizes:
            write_passage_to_word(passage.get("title"), passage.get("text"), passage.get("grade"), size, dir_path)

def write_passage_to_word(title, body, grade, size, dir_path):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(size)

    document.add_paragraph((title + "\n" + body), style='Normal')

    sub_dir = dir_path + grade
    if not os.path.exists(sub_dir):
        os.makedirs(sub_dir)
    document.save((sub_dir + "/" + grade + "_size" + str(size) + ".docx"))

def write_sight_words(form, dir_path, sizes):
    pre_primer = form.get("pre_primer")
    primer = form.get("primer")
    grade_1 = form.get("grade_1")
    grade_2 = form.get("grade_2")
    grade_3 = form.get("grade_3")
    grade_4 = form.get("grade_4")
    grade_5 = form.get("grade_5")
    grade_6 = form.get("grade_6")
    grade_7 = form.get("grade_7")
    grade_8 = form.get("grade_8")
    grade_9 = form.get("grade_9")
    grade_10 = form.get("grade_10")
    grade_11 = form.get("grade_11")
    grade_12 = form.get("grade_12")

    pre_primer_string = "\n".join(pre_primer)
    primer_string = "\n".join(primer)
    grade_1_string = "\n".join(grade_1)
    grade_2_string = "\n".join(grade_2)
    grade_3_string = "\n".join(grade_3)
    grade_4_string = "\n".join(grade_4)
    grade_5_string = "\n".join(grade_5)
    grade_6_string = "\n".join(grade_6)
    grade_7_string = "\n".join(grade_7)
    grade_8_string = "\n".join(grade_8)
    grade_9_string = "\n".join(grade_9)
    grade_10_string = "\n".join(grade_10)
    grade_11_string = "\n".join(grade_11)
    grade_12_string = "\n".join(grade_12)

    for size in sizes:
        write_string_to_word(pre_primer_string, size, dir_path, "pre_primer")
        write_string_to_word(primer_string, size, dir_path, "primer")
        write_string_to_word(grade_1_string, size, dir_path, "grade_1")
        write_string_to_word(grade_2_string, size, dir_path, "grade_2")
        write_string_to_word(grade_3_string, size, dir_path, "grade_3")
        write_string_to_word(grade_4_string, size, dir_path, "grade_4")
        write_string_to_word(grade_5_string, size, dir_path, "grade_5")
        write_string_to_word(grade_6_string, size, dir_path, "grade_6")
        write_string_to_word(grade_7_string, size, dir_path, "grade_7")
        write_string_to_word(grade_8_string, size, dir_path, "grade_8")
        write_string_to_word(grade_9_string, size, dir_path, "grade_9")
        write_string_to_word(grade_10_string, size, dir_path, "grade_10")
        write_string_to_word(grade_11_string, size, dir_path, "grade_11")
        write_string_to_word(grade_12_string, size, dir_path, "grade_12")

def write_string_to_word(text_string, size, dir_path, name):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(size)

    document.add_paragraph(text_string, style='Normal')
    sub_dir = dir_path + name
    if not os.path.exists(sub_dir):
        os.makedirs(sub_dir)
    document.save((sub_dir + "/" + name + "_size" + str(size) + ".docx"))

if __name__ == '__main__':
    main()